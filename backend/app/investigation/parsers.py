from __future__ import annotations

import csv
import hashlib
import xml.etree.ElementTree as ET
from io import BytesIO, StringIO
from pathlib import Path

from pypdf import PdfReader

from app.investigation.models import ContentChunk, DocumentType, ParsedDocument


def detect_document_type(path: Path) -> DocumentType:
    suffix = path.suffix.lower()
    mapping: dict[str, DocumentType] = {
        ".txt": DocumentType.TXT,
        ".pdf": DocumentType.PDF,
        ".xlsx": DocumentType.XLSX,
        ".csv": DocumentType.CSV,
        ".docx": DocumentType.DOCX,
        ".xml": DocumentType.XML,
    }
    return mapping.get(suffix, DocumentType.UNKNOWN)


def generate_doc_id(path: Path) -> str:
    return hashlib.sha256(str(path.resolve()).encode()).hexdigest()[:16]


def parse_txt(path: Path) -> ParsedDocument:
    """Parse tab/semicolon-separated GDPdU accounting exports."""
    raw = path.read_bytes()
    # Try common encodings for German accounting data
    for encoding in ("utf-8", "cp1252", "iso-8859-1"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode("utf-8", errors="replace")

    lines = text.strip().splitlines()
    chunks: list[ContentChunk] = []

    # Detect delimiter: semicolon is used in GDPdU TXT exports
    delimiter = ";"
    if lines and "\t" in lines[0] and ";" not in lines[0]:
        delimiter = "\t"

    for row_num, line in enumerate(lines, start=1):
        stripped = line.strip()
        if not stripped:
            continue
        chunks.append(ContentChunk(
            text=stripped,
            source_ref=f"{path.name}:row:{row_num}",
            chunk_index=row_num - 1,
        ))

    return ParsedDocument(
        doc_id=generate_doc_id(path),
        filename=path.name,
        doc_type=DocumentType.TXT,
        content_chunks=chunks,
        metadata={"delimiter": delimiter, "row_count": str(len(chunks))},
        page_count=0,
    )


def parse_gdpdu_index(path: Path) -> ParsedDocument:
    """Parse GDPdU index.xml to extract column schemas and table relationships."""
    raw = path.read_bytes()
    tree = ET.parse(BytesIO(raw))
    root = tree.getroot()

    chunks: list[ContentChunk] = []
    metadata: dict[str, str] = {}

    # Extract data supplier info
    supplier = root.find("DataSupplier")
    if supplier is not None:
        name_el = supplier.find("Name")
        if name_el is not None and name_el.text:
            metadata["supplier_name"] = name_el.text
        location_el = supplier.find("Location")
        if location_el is not None and location_el.text:
            metadata["supplier_location"] = location_el.text
        comment_el = supplier.find("Comment")
        if comment_el is not None and comment_el.text:
            metadata["comment"] = comment_el.text

    # Extract table schemas
    chunk_idx = 0
    tables: list[str] = []
    for media in root.iter("Media"):
        for table in media.iter("Table"):
            table_url = table.find("URL")
            table_name = table.find("Name")
            table_desc = table.find("Description")

            url_text = table_url.text if table_url is not None and table_url.text else ""
            name_text = table_name.text if table_name is not None and table_name.text else url_text
            desc_text = table_desc.text if table_desc is not None and table_desc.text else ""

            columns: list[str] = []
            var_length = table.find("VariableLength")
            if var_length is not None:
                for col in var_length.iter("VariableColumn"):
                    col_name = col.find("Name")
                    if col_name is not None and col_name.text:
                        columns.append(col_name.text)

            schema_text = (
                f"Table: {name_text}\n"
                f"File: {url_text}\n"
                f"Description: {desc_text}\n"
                f"Columns: {', '.join(columns)}"
            )
            chunks.append(ContentChunk(
                text=schema_text,
                source_ref=f"{path.name}:table:{name_text}",
                chunk_index=chunk_idx,
            ))
            chunk_idx += 1
            tables.append(name_text)

    metadata["tables"] = ";".join(tables)

    return ParsedDocument(
        doc_id=generate_doc_id(path),
        filename=path.name,
        doc_type=DocumentType.XML,
        content_chunks=chunks,
        metadata=metadata,
        page_count=0,
    )


def parse_pdf(path: Path) -> ParsedDocument:
    """Parse PDF using pypdf, preserving page numbers."""
    file_bytes = path.read_bytes()
    reader = PdfReader(BytesIO(file_bytes))
    chunks: list[ContentChunk] = []

    for page_num, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            chunks.append(ContentChunk(
                text=text,
                source_ref=f"{path.name}:page:{page_num}",
                chunk_index=page_num - 1,
            ))

    return ParsedDocument(
        doc_id=generate_doc_id(path),
        filename=path.name,
        doc_type=DocumentType.PDF,
        content_chunks=chunks,
        metadata={},
        page_count=len(reader.pages),
    )


def parse_xlsx(path: Path) -> ParsedDocument:
    """Parse Excel .xlsx using openpyxl, converting sheets to row content."""
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    chunks: list[ContentChunk] = []
    chunk_idx = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row_num, row in enumerate(ws.iter_rows(values_only=True), start=1):
            cell_values = [str(cell) if cell is not None else "" for cell in row]
            row_text = ";".join(cell_values).strip(";").strip()
            if row_text and row_text != ";" * len(cell_values):
                chunks.append(ContentChunk(
                    text=row_text,
                    source_ref=f"{path.name}:{sheet_name}:row:{row_num}",
                    chunk_index=chunk_idx,
                ))
                chunk_idx += 1

    wb.close()

    return ParsedDocument(
        doc_id=generate_doc_id(path),
        filename=path.name,
        doc_type=DocumentType.XLSX,
        content_chunks=chunks,
        metadata={"sheet_count": str(len(wb.sheetnames))},
        page_count=0,
    )


def parse_csv(path: Path) -> ParsedDocument:
    """Parse CSV with auto-detected delimiter, preserving row numbers."""
    raw = path.read_bytes()
    for encoding in ("utf-8", "cp1252", "iso-8859-1"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode("utf-8", errors="replace")

    # Auto-detect delimiter
    sniffer = csv.Sniffer()
    sample = text[:4096]
    try:
        dialect = sniffer.sniff(sample, delimiters=";,\t|")
        delimiter = dialect.delimiter
    except csv.Error:
        delimiter = ";" if ";" in sample else ","

    chunks: list[ContentChunk] = []
    reader = csv.reader(StringIO(text), delimiter=delimiter)

    for row_num, row in enumerate(reader, start=1):
        row_text = delimiter.join(row).strip()
        if row_text:
            chunks.append(ContentChunk(
                text=row_text,
                source_ref=f"{path.name}:row:{row_num}",
                chunk_index=row_num - 1,
            ))

    return ParsedDocument(
        doc_id=generate_doc_id(path),
        filename=path.name,
        doc_type=DocumentType.CSV,
        content_chunks=chunks,
        metadata={"delimiter": delimiter, "row_count": str(len(chunks))},
        page_count=0,
    )


def parse_docx(path: Path) -> ParsedDocument:
    """Parse DOCX using python-docx, extracting paragraph text."""
    from docx import Document

    doc = Document(str(path))
    chunks: list[ContentChunk] = []
    chunk_idx = 0

    for para_num, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if text:
            chunks.append(ContentChunk(
                text=text,
                source_ref=f"{path.name}:paragraph:{para_num}",
                chunk_index=chunk_idx,
            ))
            chunk_idx += 1

    # Also extract tables
    for table_num, table in enumerate(doc.tables, start=1):
        for row_num, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells)
            if row_text.strip(" |"):
                chunks.append(ContentChunk(
                    text=row_text,
                    source_ref=f"{path.name}:table:{table_num}:row:{row_num}",
                    chunk_index=chunk_idx,
                ))
                chunk_idx += 1

    return ParsedDocument(
        doc_id=generate_doc_id(path),
        filename=path.name,
        doc_type=DocumentType.DOCX,
        content_chunks=chunks,
        metadata={"paragraph_count": str(len(doc.paragraphs))},
        page_count=0,
    )


PARSER_DISPATCH: dict[DocumentType, type[None] | None] = {}


def parse_document(path: Path) -> ParsedDocument:
    """Dispatch to the correct parser based on file extension."""
    doc_type = detect_document_type(path)

    # Special case: index.xml files are GDPdU schema definitions
    if path.name.lower() == "index.xml":
        return parse_gdpdu_index(path)

    parsers = {
        DocumentType.TXT: parse_txt,
        DocumentType.PDF: parse_pdf,
        DocumentType.XLSX: parse_xlsx,
        DocumentType.CSV: parse_csv,
        DocumentType.DOCX: parse_docx,
        DocumentType.XML: parse_gdpdu_index,
    }

    parser = parsers.get(doc_type)
    if parser is None:
        # Return a minimal document for unsupported types
        return ParsedDocument(
            doc_id=generate_doc_id(path),
            filename=path.name,
            doc_type=DocumentType.UNKNOWN,
            content_chunks=[],
            metadata={"error": f"No parser for extension: {path.suffix}"},
            page_count=0,
        )

    return parser(path)
